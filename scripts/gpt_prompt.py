import openai
from settings import OPENAI_API_KEY
from docx import Document
import os

openai.api_key = OPENAI_API_KEY

def gen_text(resume, job_listing, your_name, hiring_manager):
    job_listing_response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role":"system", "content": "You are a entry level data scientist creating cover letters."},
            {"role":"user", "content": f"Summarize the companies objectives, qualifications and requirements from the "
                                       f"following job post and it is relevent to information found in a cover letter. "
                                       f"{job_listing}"}],
        temperature=0,
        max_tokens=2048
    )
    jl_result = job_listing_response['choices'][0]['message']['content']

    background_response = openai.ChatCompletion.create(
        model='gpt-3.5-turbo',
        messages=[
            {"role":"assistant", "content": f"Professionally identify relationships "
                                            f"between items on the following resume and job listing summary. Prioritize the most relevant items first RESUME: "
                                            f"{resume}, JOB LISTING SUMMARY: {jl_result}"}
        ],
        temperature=0,
        max_tokens=2048
    )
    background_result = background_response['choices'][0]['message']['content']

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user",
             "content": f"Use at most 250 words and don't be repetative. Using the following summary of a job listing and background create a "
                        f"cover letter using a 5th "
                        f"grade writing level. BACKGROUND: {background_result}, JOB LISTING: {jl_result}. The applicants "
                        f"name is {your_name} and the hiring managers name is {hiring_manager}."}],
        temperature=0,
        max_tokens=2048
    )
    result = response['choices'][0]['message']['content']
    return result

def obtainText(docFileName):
    document = Document(docFileName)
    finalText = []
    for line in document.paragraphs:
        finalText.append(line.text)

    # read table data
    if len(document.tables) > 0:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    finalText.append(cell.text)

    return '\n'.join(finalText)

def saveText(textData):
    document = Document()


# resume = obtainText(r'C:\Users\kdenn\PycharmProjects\cover_letter_generator\static\Kevin Resume_072223.docx')
# job_listing = obtainText(r'C:\Users\kdenn\PycharmProjects\cover_letter_generator\static\job_listing.docx')
#
# #print(response['choices'][0]['text'])
# print(gen_text(resume, job_listing))